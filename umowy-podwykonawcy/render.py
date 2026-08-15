#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Renderer: lista bloków treści -> .docx i .md.

Typy bloków:
    h1    tytuł dokumentu
    sub   podtytuł (kursywa, wyśrodkowany)
    h2    nagłówek paragrafu (resetuje numerację ustępów)
    p     zwykły akapit
    li    ustęp numerowany (1., 2., …)
    lia   litera (a), b), …) — numeracja resetuje się przy każdym li/p/h2
    note  akapit mniejszą czcionką, kursywą (komentarz redakcyjny)
    table (['nagłówek','nagłówek'], [['a','b'], …])
    sig   blok podpisów
    break twarda pauza strony

W tekście działa **pogrubienie**.
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

LETTERS = "abcdefghijklmnopqrstuvwxyz"


def fill(text, placeholders):
    for k, v in placeholders.items():
        text = text.replace(k, v)
    return text


def _runs(par, text):
    for chunk in re.split(r"(\*\*.+?\*\*)", text):
        if not chunk:
            continue
        run = par.add_run(chunk[2:-2] if chunk.startswith("**") else chunk)
        run.bold = chunk.startswith("**")


def build_docx(blocks, placeholders, path, sig_left="Zleceniodawca", sig_right="Podwykonawca"):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.2)

    n_li = n_lia = 0
    for kind, raw in blocks:
        text = fill(raw, placeholders) if isinstance(raw, str) else raw

        if kind == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(14)
        elif kind == "sub":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(14)
            p.add_run(text).italic = True
        elif kind == "h2":
            n_li = n_lia = 0
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            p.add_run(text).bold = True
        elif kind == "p":
            n_lia = 0
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _runs(p, text)
        elif kind == "note":
            n_lia = 0
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            r.italic = True
            r.font.size = Pt(9)
        elif kind == "li":
            n_li += 1
            n_lia = 0
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)
            p.add_run(f"{n_li}. ")
            _runs(p, text)
        elif kind == "lia":
            letter = LETTERS[n_lia]
            n_lia += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            p.add_run(f"{letter}) ")
            _runs(p, text)
        elif kind == "table":
            header, rows = text
            t = doc.add_table(rows=1 + len(rows), cols=len(header))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(header):
                cell = t.rows[0].cells[i]
                cell.text = ""
                cell.paragraphs[0].add_run(h).bold = True
            for ri, row in enumerate(rows, start=1):
                for ci, val in enumerate(row):
                    c = t.rows[ri].cells[ci]
                    c.text = ""
                    _runs(c.paragraphs[0], fill(val, placeholders))
            for row in t.rows:
                for c in row.cells:
                    for par in c.paragraphs:
                        par.paragraph_format.space_after = Pt(2)
                        for r in par.runs:
                            r.font.size = Pt(9.5)
            doc.add_paragraph()
        elif kind == "break":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif kind == "sig":
            doc.add_paragraph().paragraph_format.space_after = Pt(36)
            t = doc.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = sig_left
            t.rows[0].cells[1].text = sig_right
            t.rows[1].cells[0].text = "\n\n……………………………………………"
            t.rows[1].cells[1].text = "\n\n……………………………………………"
            for row in t.rows:
                for c in row.cells:
                    for par in c.paragraphs:
                        par.paragraph_format.space_after = Pt(0)

    doc.save(path)


def build_md(blocks, placeholders, path, sig_left="Zleceniodawca", sig_right="Podwykonawca"):
    out, n_li, n_lia = [], 0, 0
    for kind, raw in blocks:
        text = fill(raw, placeholders) if isinstance(raw, str) else raw
        if kind == "h1":
            out.append(f"# {text}\n")
        elif kind == "sub":
            out.append(f"*{text}*\n")
        elif kind == "h2":
            n_li = n_lia = 0
            out.append(f"\n## {text}\n")
        elif kind in ("p", "note"):
            n_lia = 0
            out.append((f"*{text}*\n" if kind == "note" else f"{text}\n"))
        elif kind == "li":
            n_li += 1
            n_lia = 0
            out.append(f"{n_li}. {text}\n")
        elif kind == "lia":
            out.append(f"    {LETTERS[n_lia]}) {text}\n")
            n_lia += 1
        elif kind == "table":
            header, rows = text
            out.append("| " + " | ".join(header) + " |")
            out.append("|" + "---|" * len(header))
            for row in rows:
                out.append("| " + " | ".join(fill(v, placeholders).replace("\n", " ") for v in row) + " |")
            out.append("")
        elif kind == "break":
            out.append("\n---\n")
        elif kind == "sig":
            out.append(f"\n| {sig_left} | {sig_right} |\n|---|---|\n"
                       "| …………………………… | …………………………… |\n")
    open(path, "w", encoding="utf-8").write("\n".join(out))
